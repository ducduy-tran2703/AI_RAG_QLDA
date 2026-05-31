"""
extract_docx_optimized.py
==========================
Module bóc tách TỐI ƯU cho LLM - Văn bản hành chính (.docx)
Dự án: AI_RAG_QLDA - Kiểm tra thể thức theo NĐ30/2020

🎯 ĐẶC ĐIỂM CHÍNH:
──────────────────
✅ Bước 1: Tái cấu trúc Block + Inline (Run-level)
   - Block: Thuộc tính đoạn văn (căn lề, thụt dòng, khoảng cách)
   - Runs: Gom các ký tự liền kề có CÙNG định dạng thành cụm

✅ Bước 2: Lọc rác (Drop Nulls & Defaults)
   - Loại bỏ hoàn toàn: None, False, [], 0, giá trị mặc định
   - Chỉ giữ thông tin CÓ GIÁ TRỊ cho suy luận logic
   - Không lưu tọa độ bbox (dành cho PDF)

📦 OUTPUT CÔ ĐẶC:
────────────────
{
  "filename": "cv123.docx",
  "margins": {"left_cm": 3.0, "right_cm": 2.0, ...},
  "paragraphs": [
    {
      "idx": 0,
      "text": "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
      "block": {
        "align": "CENTER",
        "space_before_pt": 12
      },
      "runs": [
        {"fmt": "Arial-13pt-Bold", "txt": "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"}
      ],
      "zone": "RIGHT_COL"
    },
    {
      "idx": 1,
      "text": "Điều 1. Phạm vi điều chỉnh",
      "block": {
        "indent_cm": 1.0,
        "space_before_pt": 6
      },
      "runs": [
        {"fmt": "14pt-Bold", "txt": "Điều 1."},
        {"fmt": "14pt", "txt": " Phạm vi điều chỉnh"}
      ],
      "zone": "FULL_WIDTH"
    }
  ]
}

Token giảm ~70-85% so với bản cũ! 🚀
"""

import json
import zipfile
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import re as _re_top  # dùng cho numbered list


# ═════════════════════════════════════════════════════════════════════
# HẰNG SỐ QUY ĐỔI ĐƠN VỊ
# ═════════════════════════════════════════════════════════════════════
EMU_PER_CM    = 914400 / 2.54   # 1 cm = 360000 EMU
TWIP_PER_CM   = 1440 / 2.54     # 1 cm ≈ 566.93 twip
PT_PER_TWIP   = 1 / 20          # 1 twip = 1/20 pt


def emu_to_cm(emu: Optional[int]) -> Optional[float]:
    if emu is None:
        return None
    return round(emu / EMU_PER_CM, 2)


def twip_to_cm(twip: Optional[int]) -> Optional[float]:
    if twip is None:
        return None
    return round(twip / TWIP_PER_CM, 2)


def twip_to_pt(twip: Optional[int]) -> Optional[float]:
    if twip is None:
        return None
    return round(twip * PT_PER_TWIP, 2)


def emu_to_pt(emu: Optional[int]) -> Optional[float]:
    if emu is None:
        return None
    return round(emu / 12700, 2)


# ═════════════════════════════════════════════════════════════════════
# HELPER: Phát hiện loại văn bản (doc_type)
# ═════════════════════════════════════════════════════════════════════
# Ánh xạ từ khóa → loại văn bản theo NĐ30
_DOC_TYPE_PATTERNS: List[tuple] = [
    # (keyword_lowercase, doc_type)
    # ĐẶT từ khóa dài/cụ thể TRƯỚC để tránh match nhầm
    ("văn bản hợp nhất", "vbhn"),
    ("thông tư liên tịch", "thong_tu_lien_tich"),
    ("nghị định",          "nghi_dinh"),
    ("thông tư",           "thong_tu"),
    ("quyết định",         "quyet_dinh"),
    ("chỉ thị",            "chi_thi"),
    ("nghị quyết",         "nghi_quyet"),
    ("công văn",           "cong_van"),
    ("tờ trình",           "to_trinh"),
    ("báo cáo",            "bao_cao"),
    ("kế hoạch",           "ke_hoach"),
    ("hướng dẫn",          "huong_dan"),
    ("thông báo",          "thong_bao"),
    ("biên bản",           "bien_ban"),
    ("hợp đồng",           "hop_dong"),
    ("giấy mời",           "giay_moi"),
    ("giấy phép",          "giay_phep"),
    ("tờ khai",            "to_khai"),
    ("đề án",              "de_an"),
    ("chương trình",       "chuong_trinh"),
]

# Detect từ tên file — ưu tiên cao nhất vì tên file do người dùng đặt có chủ đích
_FILENAME_PATTERNS: List[tuple] = [
    # (pattern_in_filename_lowercase, doc_type)
    ("vbhn",   "vbhn"),
    ("-nd-",   "nghi_dinh"),
    ("/nd-",   "nghi_dinh"),
    ("-tt-",   "thong_tu"),
    ("/tt-",   "thong_tu"),
    ("-ttlt-", "thong_tu_lien_tich"),
    ("-qd-",   "quyet_dinh"),
    ("/qd-",   "quyet_dinh"),
    ("-ct-",   "chi_thi"),
    ("-nq-",   "nghi_quyet"),
    ("-cv-",   "cong_van"),
]


def _detect_doc_type_from_filename(filename: str) -> Optional[str]:
    """
    Phát hiện loại văn bản từ tên file.
    VD: '23-vbhn-vpqh.docx' → 'vbhn', '45-2013-nd-cp.docx' → 'nghi_dinh'
    Trả về doc_type hoặc None nếu không nhận ra.
    """
    name_lower = filename.lower()
    for pattern, doc_type in _FILENAME_PATTERNS:
        if pattern in name_lower:
            return doc_type
    return None


def _detect_doc_type(paragraphs_text: List[str], filename: str = "") -> str:
    """
    Phát hiện loại văn bản.
    Ưu tiên: tên file → nội dung 20 đoạn đầu → 'unknown'.
    """
    # 1. Ưu tiên filename trước (chắc chắn nhất)
    if filename:
        doc_type = _detect_doc_type_from_filename(filename)
        if doc_type:
            return doc_type

    # 2. Fallback: dò nội dung văn bản
    sample = " ".join(paragraphs_text[:20]).lower()
    for keyword, doc_type in _DOC_TYPE_PATTERNS:
        if keyword in sample:
            return doc_type

    return "unknown"


# ═════════════════════════════════════════════════════════════════════
# HELPERS: Theme Fonts & Style Map
# ═════════════════════════════════════════════════════════════════════
def _get_theme_fonts(docx_path: str) -> dict:
    """Bóc tách theme fonts (major/minor) từ theme1.xml"""
    result = {'major': None, 'minor': None}
    try:
        with zipfile.ZipFile(docx_path) as z:
            if 'word/theme/theme1.xml' not in z.namelist():
                return result
            tree = etree.fromstring(z.read('word/theme/theme1.xml'))
            ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
            for el in tree.findall('.//a:majorFont/a:latin', ns):
                result['major'] = el.get('typeface')
            for el in tree.findall('.//a:minorFont/a:latin', ns):
                result['minor'] = el.get('typeface')
    except Exception:
        pass
    return result


def _get_numbering_map(docx_path: str) -> dict:
    """
    Đọc numbering.xml, trả về map:
      (numId, ilvl) → start_value (int)

    Dùng để bóc số thứ tự tự động (Word List Numbering) vào text.
    VD: numId=1, ilvl=0, start=1 → mục "1.", "2.", "3."...
    """
    result = {}
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    wval = f'{{{W}}}val'
    try:
        with zipfile.ZipFile(docx_path) as z:
            if 'word/numbering.xml' not in z.namelist():
                return result
            tree = etree.fromstring(z.read('word/numbering.xml'))
            ns = {'w': W}

            # abstractNum: lưu ilvl → start
            abstract_starts = {}
            for ab in tree.findall('w:abstractNum', ns):
                ab_id = ab.get(f'{{{W}}}abstractNumId')
                ab_starts = {}
                for lvl in ab.findall('w:lvl', ns):
                    ilvl = lvl.get(f'{{{W}}}ilvl', '0')
                    start_el = lvl.find('w:start', ns)
                    numFmt_el = lvl.find('w:numFmt', ns)
                    fmt = numFmt_el.get(wval, '') if numFmt_el is not None else ''
                    start = int(start_el.get(wval, 1)) if start_el is not None else 1
                    ab_starts[ilvl] = {'start': start, 'fmt': fmt}
                abstract_starts[ab_id] = ab_starts

            # num: numId → abstractNumId (+ xử lý lvlOverride/startOverride)
            for num_el in tree.findall('w:num', ns):
                num_id = num_el.get(f'{{{W}}}numId')
                ref = num_el.find('w:abstractNumId', ns)
                if ref is None:
                    continue
                ab_id = ref.get(wval)
                ab_data = abstract_starts.get(ab_id, {})
                for ilvl, data in ab_data.items():
                    result[(num_id, ilvl)] = dict(data)  # copy để tránh mutate abstractNum gốc

                # Đọc lvlOverride: Word dùng khi người dùng "Continue numbering"
                # hoặc "Set numbering value" → startOverride ghi đè start của abstractNum
                for ov in num_el.findall('w:lvlOverride', ns):
                    olvl = ov.get(f'{{{W}}}ilvl', '0')
                    startOverride_el = ov.find('w:startOverride', ns)
                    if startOverride_el is not None:
                        override_val = int(startOverride_el.get(wval, 1))
                        key = (num_id, olvl)
                        if key in result:
                            result[key] = {**result[key], 'start': override_val}
                        else:
                            result[key] = {'start': override_val, 'fmt': 'decimal'}
    except Exception:
        pass
    return result


def _get_para_numpr(para) -> tuple:
    """
    Lấy (numId, ilvl) từ paragraph nếu nó thuộc numbered list.
    Trả về (None, None) nếu không phải list.
    """
    ppr = para._p.find(qn('w:pPr'))
    if ppr is None:
        return None, None
    numPr = ppr.find(qn('w:numPr'))
    if numPr is None:
        return None, None
    numId_el = numPr.find(qn('w:numId'))
    ilvl_el  = numPr.find(qn('w:ilvl'))
    if numId_el is None:
        return None, None
    numId = numId_el.get(qn('w:val'))
    ilvl  = ilvl_el.get(qn('w:val'), '0') if ilvl_el is not None else '0'
    # numId=0 nghĩa là tắt list
    if numId == '0':
        return None, None
    return numId, ilvl


def _get_style_fonts_from_zip(docx_path: str) -> dict:
    """Bóc tách font + size từ styles.xml (fallback cho runs không có format trực tiếp)"""
    result = {}
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    wval = f'{{{W}}}val'
    try:
        with zipfile.ZipFile(docx_path) as z:
            styles_xml = etree.fromstring(z.read('word/styles.xml'))
            ns = {'w': W}
            for style_el in styles_xml.findall('.//w:style', ns):
                name_el = style_el.find('w:name', ns)
                if name_el is None:
                    continue
                sname = name_el.get(wval, '')
                rfonts = style_el.find('.//w:rFonts', ns)
                sz_el  = style_el.find('.//w:sz', ns)
                font = None
                size = None
                if rfonts is not None:
                    font = (rfonts.get(f'{{{W}}}ascii')
                            or rfonts.get(f'{{{W}}}hAnsi'))
                if sz_el is not None:
                    size = int(sz_el.get(wval, 0)) / 2
                result[sname] = {'font': font, 'size': size}
    except Exception:
        pass
    return result


# ═════════════════════════════════════════════════════════════════════
# HELPER: Số trang (Header/Footer)
# ═════════════════════════════════════════════════════════════════════
def _extract_page_number_info(doc: Document, docx_path: str) -> Optional[Dict[str, Any]]:
    """
    Bóc tách thông tin số trang từ header/footer.
    Chỉ trả về dict nếu tìm thấy số trang, None nếu không có.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ALIGN_MAP = {'left': 'LEFT', 'right': 'RIGHT', 'center': 'CENTER', 'both': 'JUSTIFY'}

    found         = False
    font_size_pt  = None
    alignment     = None
    position      = None

    # Kiểm tra "Different First Page" → hidden_on_page1
    sec = doc.sections[0]
    sec_pr = sec._sectPr
    title_pg = sec_pr.find(qn('w:titlePg'))
    different_first = title_pg is not None

    def _scan_header_footer(xml_bytes: bytes, pos: str):
        nonlocal found, font_size_pt, alignment, position
        try:
            tree = etree.fromstring(xml_bytes)
            ns = {'w': W}
            for instr in tree.findall('.//w:instrText', ns):
                if 'PAGE' in (instr.text or ''):
                    found = True
                    position = pos
                    para_el = instr
                    while para_el is not None and para_el.tag != f'{{{W}}}p':
                        para_el = para_el.getparent()
                    if para_el is not None:
                        jc = para_el.find(f'.//{{{W}}}jc')
                        if jc is not None:
                            alignment = ALIGN_MAP.get(jc.get(f'{{{W}}}val', ''), None)
                        sz = para_el.find(f'.//{{{W}}}sz')
                        if sz is not None:
                            val = sz.get(f'{{{W}}}val')
                            if val:
                                font_size_pt = int(val) / 2
                    break
            for fld in tree.findall('.//w:fldSimple', ns):
                instr_attr = fld.get(f'{{{W}}}instr', '')
                if 'PAGE' in instr_attr:
                    found = True
                    position = pos
                    para_el = fld
                    while para_el is not None and para_el.tag != f'{{{W}}}p':
                        para_el = para_el.getparent()
                    if para_el is not None:
                        jc = para_el.find(f'.//{{{W}}}jc')
                        if jc is not None:
                            alignment = ALIGN_MAP.get(jc.get(f'{{{W}}}val', ''), None)
                        sz = para_el.find(f'.//{{{W}}}sz')
                        if sz is not None:
                            val = sz.get(f'{{{W}}}val')
                            if val:
                                font_size_pt = int(val) / 2
                    break
        except Exception:
            pass

    try:
        with zipfile.ZipFile(docx_path) as z:
            names = z.namelist()
            for name in names:
                if 'word/header' in name and name.endswith('.xml'):
                    _scan_header_footer(z.read(name), 'HEADER')
                elif 'word/footer' in name and name.endswith('.xml'):
                    _scan_header_footer(z.read(name), 'FOOTER')
    except Exception:
        pass

    if not found:
        return None

    # Lọc rác: chỉ giữ thông tin có giá trị
    result = {}
    if font_size_pt:
        result['size_pt'] = font_size_pt
    if alignment:
        result['align'] = alignment
    if position:
        result['pos'] = position
    if different_first:
        result['hidden_p1'] = True

    return result if result else None


# ═════════════════════════════════════════════════════════════════════
# HELPER: Paragraph Border Bottom
# ═════════════════════════════════════════════════════════════════════
def _extract_border(para) -> Optional[Dict[str, Any]]:
    """Bóc tách đường kẻ ngang (pBdr/bottom) dưới đoạn văn."""
    ppr = para._p.find(qn('w:pPr'))
    if ppr is None:
        return None
    pbdr = ppr.find(qn('w:pBdr'))
    if pbdr is None:
        return None
    bottom = pbdr.find(qn('w:bottom'))
    if bottom is None:
        return None

    bstyle = bottom.get(qn('w:val'), None)
    if bstyle in (None, 'none', 'nil'):
        return None

    # Lọc rác: chỉ lưu thông tin cần thiết
    result = {}
    color   = bottom.get(qn('w:color'), None)
    sz_val  = bottom.get(qn('w:sz'), None)
    if color:
        result['color'] = color
    if sz_val:
        result['width_pt'] = round(int(sz_val) / 8, 2)

    return result if result else {'has': True}


# ═════════════════════════════════════════════════════════════════════
# HELPER: Spatial Position
# ═════════════════════════════════════════════════════════════════════
def _infer_spatial_position(para, page_width_cm: float,
                             left_margin_cm: float, right_margin_cm: float) -> Optional[str]:
    """
    Phân loại vị trí không gian:
      - LEFT_COL  : cột trái (cơ quan, số ký hiệu, nơi nhận)
      - RIGHT_COL : cột phải (Quốc hiệu, địa danh, chức vụ)
      - FULL_WIDTH: toàn bộ chiều ngang (nội dung chính)
    """
    ppr = para._p.find(qn('w:pPr'))
    if ppr is None:
        return None

    jc = ppr.find(qn('w:jc'))
    if jc is not None:
        val = jc.get(qn('w:val'), '')
        if val == 'right':
            return 'RIGHT_COL'
        if val == 'center':
            # Canh giữa + indent trái lớn → cột phải
            ind = ppr.find(qn('w:ind'))
            if ind is not None:
                left_twip = int(ind.get(qn('w:left'), 0) or 0)
                content_width = page_width_cm - left_margin_cm - right_margin_cm
                if twip_to_cm(left_twip) and twip_to_cm(left_twip) > content_width * 0.4:
                    return 'RIGHT_COL'
            return None  # CENTER không phải zone đặc biệt

    return None  # LEFT/JUSTIFY → FULL_WIDTH (mặc định)


# ═════════════════════════════════════════════════════════════════════
# CORE: Gom cụm định dạng (Run-level Accumulation)
# ═════════════════════════════════════════════════════════════════════
@dataclass
class RunInfo:
    """Cụm text có cùng định dạng (Font, Size, Bold, Italic)"""
    fmt: str        # "Times-14pt-Bold" hoặc "14pt-Bold" nếu không có font
    txt: str        # Nội dung text


def _extract_runs(para, style_map: dict, theme_fonts: dict) -> List[RunInfo]:
    """
    Gom các Run có cùng định dạng thành cụm.
    
    QUAN TRỌNG: Thuật toán "Accumulation" - tránh tạo quá nhiều cụm nhỏ lẻ.
    """
    runs_data = []
    
    for run in para.runs:
        if not run.text:  # Bỏ qua run rỗng
            continue
        
        # ── Font name ──────────────────────────────
        font_name = None
        rpr = run._r.find(qn('w:rPr'))
        if rpr is not None:
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is not None:
                f = (rfonts.get(qn('w:ascii'))
                     or rfonts.get(qn('w:hAnsi'))
                     or rfonts.get(qn('w:cs')))
                if f and not f.startswith('+'):
                    font_name = f
                else:
                    # Theme font reference
                    theme_ref = rfonts.get(qn('w:asciiTheme')) or rfonts.get(qn('w:hAnsiTheme'))
                    if theme_ref:
                        if 'major' in theme_ref.lower():
                            font_name = theme_fonts.get('major')
                        else:
                            font_name = theme_fonts.get('minor')
        
        # Fallback từ style
        if not font_name:
            sinfo = style_map.get(para.style.name, {})
            font_name = sinfo.get('font')
        
        # ── Font size ──────────────────────────────
        font_size = None
        if rpr is not None:
            sz = rpr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                if val:
                    font_size = int(val) / 2
        
        if font_size is None:
            sinfo = style_map.get(para.style.name, {})
            font_size = sinfo.get('size')
        
        # ── Bold / Italic ──────────────────────────
        bold = run.bold
        italic = run.italic
        
        # ── Tạo format string cô đặc ───────────────
        parts = []
        if font_name:
            parts.append(font_name)
        if font_size:
            parts.append(f"{int(font_size)}pt")
        if bold:
            parts.append("Bold")
        if italic:
            parts.append("Italic")
        
        fmt = "-".join(parts) if parts else "default"
        
        runs_data.append({'fmt': fmt, 'txt': run.text})
    
    # ── ACCUMULATION: Gom các run liền kề có cùng format ──
    if not runs_data:
        return []
    
    accumulated = []
    current = runs_data[0]
    
    for i in range(1, len(runs_data)):
        if runs_data[i]['fmt'] == current['fmt']:
            # Cùng format → gom vào cụm hiện tại
            current['txt'] += runs_data[i]['txt']
        else:
            # Khác format → lưu cụm cũ, bắt đầu cụm mới
            accumulated.append(RunInfo(fmt=current['fmt'], txt=current['txt']))
            current = runs_data[i]
    
    # Lưu cụm cuối cùng
    accumulated.append(RunInfo(fmt=current['fmt'], txt=current['txt']))
    
    return accumulated


# ═════════════════════════════════════════════════════════════════════
# CORE: Bóc tách 1 Paragraph (Block + Runs)
# ═════════════════════════════════════════════════════════════════════
def _extract_paragraph(para, idx: int,
                        style_map: dict,
                        theme_fonts: dict,
                        page_width_cm: float,
                        left_margin_cm: float,
                        right_margin_cm: float,
                        in_table: bool = False,
                        table_cell: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """
    Bóc tách 1 paragraph thành cấu trúc Block + Runs tối ưu.
    
    Trả về None nếu paragraph rỗng hoặc không có thông tin giá trị.
    """
    text = para.text.strip()
    if not text:
        return None
    
    # ══════════════════════════════════════════════════════════════════
    # BLOCK LEVEL: Thuộc tính của cả đoạn văn
    # ══════════════════════════════════════════════════════════════════
    block = {}
    
    # ── Alignment ──────────────────────────────────────────────────────
    ALIGN_MAP = {'left': 'LEFT', 'right': 'RIGHT', 'center': 'CENTER', 'both': 'JUSTIFY'}
    ppr = para._p.find(qn('w:pPr'))
    if ppr is not None:
        jc = ppr.find(qn('w:jc'))
        if jc is not None:
            val = jc.get(qn('w:val'), '')
            alignment = ALIGN_MAP.get(val, val.upper())
            if alignment and alignment != 'LEFT':  # LEFT là mặc định, bỏ qua
                block['align'] = alignment
    
    # ── Space before / after (pt) ──────────────────────────────────────
    if ppr is not None:
        spacing = ppr.find(qn('w:spacing'))
        if spacing is not None:
            before_twip = spacing.get(qn('w:before'))
            after_twip  = spacing.get(qn('w:after'))
            if before_twip:
                sp = twip_to_pt(int(before_twip))
                if sp and sp > 0:
                    block['space_before_pt'] = sp
            if after_twip:
                sp = twip_to_pt(int(after_twip))
                if sp and sp > 0:
                    block['space_after_pt'] = sp
    
    # Fallback qua python-docx API
    pf = para.paragraph_format
    if 'space_before_pt' not in block and pf.space_before is not None:
        # pf.space_before trả về object Pt, KHÔNG phải int EMU → dùng .pt
        sp = round(pf.space_before.pt, 2) if hasattr(pf.space_before, 'pt') else emu_to_pt(pf.space_before)
        if sp and sp > 0:
            block['space_before_pt'] = sp
    if 'space_after_pt' not in block and pf.space_after is not None:
        sp = round(pf.space_after.pt, 2) if hasattr(pf.space_after, 'pt') else emu_to_pt(pf.space_after)
        if sp and sp > 0:
            block['space_after_pt'] = sp
    
    # ── First-line indent (cm) ─────────────────────────────────────────
    if ppr is not None:
        ind = ppr.find(qn('w:ind'))
        if ind is not None:
            fl = ind.get(qn('w:firstLine'))
            if fl:
                try:
                    indent = twip_to_cm(int(fl))
                    if indent and indent > 0:
                        block['indent_cm'] = indent
                except (ValueError, TypeError):
                    pass
    if 'indent_cm' not in block and pf.first_line_indent is not None:
        indent = emu_to_cm(pf.first_line_indent)
        if indent and indent > 0:
            block['indent_cm'] = indent
    
    # ── Line spacing ───────────────────────────────────────────────────
    if pf.line_spacing is not None:
        ls = pf.line_spacing
        if hasattr(ls, 'pt'):
            spacing_val = round(ls.pt, 2)
        else:
            spacing_val = round(float(ls), 2)
        # Chỉ lưu nếu khác 1.0 (single spacing là mặc định)
        if spacing_val and spacing_val != 1.0:
            block['line_spacing'] = spacing_val
    
    # ── Paragraph border ───────────────────────────────────────────────
    border = _extract_border(para)
    if border:
        block['border'] = border
    
    # ══════════════════════════════════════════════════════════════════
    # RUNS LEVEL: Các cụm text có cùng định dạng
    # ══════════════════════════════════════════════════════════════════
    runs = _extract_runs(para, style_map, theme_fonts)
    
    # Nếu toàn bộ đoạn chỉ có 1 run duy nhất và format = "default" → bỏ runs, chỉ giữ text
    runs_output = None
    if len(runs) == 1 and runs[0].fmt == "default":
        runs_output = None  # Không cần lưu runs nếu chỉ có 1 cụm mặc định
    elif len(runs) > 0:
        runs_output = [{'fmt': r.fmt, 'txt': r.txt} for r in runs]
    
    # ══════════════════════════════════════════════════════════════════
    # SPATIAL ZONE: LEFT_COL / RIGHT_COL / FULL_WIDTH
    # ══════════════════════════════════════════════════════════════════
    if in_table:
        # Trong bảng: dùng table_cell để xác định zone
        zone = table_cell.replace('row', 'r').replace('col', 'c') if table_cell else None
    else:
        zone = _infer_spatial_position(para, page_width_cm, left_margin_cm, right_margin_cm)
    
    # ══════════════════════════════════════════════════════════════════
    # FINAL OUTPUT: Lọc rác - chỉ giữ thông tin có giá trị
    # ══════════════════════════════════════════════════════════════════
    result = {
        'idx': idx,
        'text': text,
    }
    
    if block:
        result['block'] = block
    if runs_output:
        result['runs'] = runs_output
    if zone:
        result['zone'] = zone
    
    return result


# ═════════════════════════════════════════════════════════════════════
# HELPER: Bóc tách paragraphs trong bảng
# ═════════════════════════════════════════════════════════════════════
def _extract_table_paragraphs(doc: Document, start_idx: int,
                               style_map: dict, theme_fonts: dict,
                               page_width_cm: float,
                               left_margin_cm: float,
                               right_margin_cm: float) -> List[Dict[str, Any]]:
    """
    Bóc tách tất cả paragraph trong bảng (thường chứa Header Zone).
    """
    paragraphs = []
    idx = start_idx
    
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_label = f"r{r_idx}c{c_idx}"
                
                for para in cell.paragraphs:
                    info = _extract_paragraph(
                        para, idx, style_map, theme_fonts,
                        page_width_cm, left_margin_cm, right_margin_cm,
                        in_table=True,
                        table_cell=cell_label,
                    )
                    if info:  # Chỉ thêm nếu không None
                        paragraphs.append(info)
                        idx += 1
    
    return paragraphs


import re as _re  # re đã import ở trên nếu có, dùng alias để an toàn

# ═════════════════════════════════════════════════════════════════════
# HELPER: Cắt vùng chữ ký / dấu mộc cuối văn bản
# ═════════════════════════════════════════════════════════════════════

_SIGNATURE_PATTERNS = _re.compile(
    r"TM\.\s*ỦY\s*BAN|TM\.\s*BỘ\s*TRƯỞNG|TM\.\s*CHÍNH\s*PHỦ"
    r"|KT\.\s*CHỦ\s*TỊCH|KT\.\s*BỘ\s*TRƯỞNG|KT\.\s*THỦ\s*TƯỚNG"
    r"|THỪA\s*ỦY\s*QUYỀN|THỪA\s*LỆNH"
    r"|TM\.\s*BAN\s*CHẤP\s*HÀNH|TM\.\s*HỘI\s*ĐỒNG"
    r"|CHỦ\s*TỊCH\s*UBND|CHỦ\s*TỊCH\s*HỘI\s*ĐỒNG",
    _re.IGNORECASE | _re.UNICODE,
)

_TINY_FONT_MAX_PT = 8.0


def _drop_signature_zone(paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Loại bỏ vùng chữ ký / dấu mộc / mã nội bộ cuối văn bản (docx).

    Chiến lược:
    1. Font < 8pt → drop ngay (mã nội bộ)
    2. Khi gặp paragraph RIGHT_COL khớp _SIGNATURE_PATTERNS
       → lưu cut_idx, drop tất cả paragraph RIGHT_COL có idx >= cut_idx
    3. LEFT_COL (Nơi nhận, v.v.) không bị drop
    """
    # Bước 1: lọc font nhỏ
    result = []
    for p in paragraphs:
        runs = p.get('runs', [])
        if runs:
            sizes = []
            for r in runs:
                m = _re.search(r'(\d+(?:\.\d+)?)pt', r.get('fmt', ''))
                if m:
                    sizes.append(float(m.group(1)))
            if sizes and max(sizes) < _TINY_FONT_MAX_PT:
                continue
        result.append(p)

    # Bước 2: tìm cut_idx — idx nhỏ nhất của paragraph RIGHT_COL khớp pattern
    cut_idx = None
    for p in result:
        if p.get('zone') == 'RIGHT_COL' and _SIGNATURE_PATTERNS.search(p.get('text', '')):
            if cut_idx is None or p['idx'] < cut_idx:
                cut_idx = p['idx']

    if cut_idx is None:
        return result

    # Bước 3: drop RIGHT_COL từ cut_idx trở đi
    return [
        p for p in result
        if not (p.get('zone') == 'RIGHT_COL' and p['idx'] >= cut_idx)
    ]


# ═════════════════════════════════════════════════════════════════════
# PUBLIC API
# ═════════════════════════════════════════════════════════════════════
def extract_docx_optimized(docx_path: str) -> Dict[str, Any]:
    """
    Bóc tách cấu trúc vật lý tối ưu cho LLM.
    
    Parameters
    ----------
    docx_path : str — đường dẫn tới file .docx
    
    Returns
    -------
    dict — Cấu trúc cô đặc (Block + Runs), đã lọc rác
    """
    path = Path(docx_path)
    doc  = Document(str(path))
    
    # ── Margins ────────────────────────────────────────────────────────
    sec = doc.sections[0]
    margins = {}
    
    left_cm   = emu_to_cm(sec.left_margin)
    right_cm  = emu_to_cm(sec.right_margin)
    top_cm    = emu_to_cm(sec.top_margin)
    bottom_cm = emu_to_cm(sec.bottom_margin)
    width_cm  = emu_to_cm(sec.page_width)
    height_cm = emu_to_cm(sec.page_height)
    
    # ── Tách page_size ra field riêng, KHÔNG lẫn vào margins ──────────
    page_size: Dict[str, float] = {}
    if width_cm:
        page_size['width_cm']  = width_cm
    if height_cm:
        page_size['height_cm'] = height_cm

    # margins chỉ còn lề (top/bottom/left/right), không có kích thước trang
    if left_cm   and left_cm   > 0: margins['left_cm']   = left_cm
    if right_cm  and right_cm  > 0: margins['right_cm']  = right_cm
    if top_cm    and top_cm    > 0: margins['top_cm']    = top_cm
    if bottom_cm and bottom_cm > 0: margins['bottom_cm'] = bottom_cm

    page_width_cm   = width_cm  or 21.0
    page_height_cm  = height_cm or 29.7
    left_margin_cm  = left_cm   or 3.0
    right_margin_cm = right_cm  or 2.0

    # ── Style map + theme fonts + numbering ───────────────────────────
    theme_fonts  = _get_theme_fonts(str(path))
    style_map    = _get_style_fonts_from_zip(str(path))
    numbering_map = _get_numbering_map(str(path))  # (numId,ilvl) → {start,fmt}
    # Bộ đếm số thứ tự hiện tại theo (numId, ilvl)
    num_counters: dict = {}

    # ── Số trang (header/footer) ───────────────────────────────────────
    page_number = _extract_page_number_info(doc, str(path))

    # ── Xác định para nằm trong bảng để skip khi duyệt thường ──────────
    # Văn bản hành chính VN hay dùng bảng ẩn để layout 2 cột ở đầu trang
    # (tên cơ quan LEFT | quốc hiệu RIGHT). Phải xử lý bảng TRƯỚC để các
    # para này có idx nhỏ → build_meta_chunk nhận ra quốc hiệu, tên cơ quan.
    table_para_ids: set = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_para_ids.add(id(para._p))

    paragraphs: List[Dict[str, Any]] = []

    # ── Bước 1: Paragraphs trong bảng TRƯỚC (header 2 cột) ────────────
    table_paras = _extract_table_paragraphs(
        doc, start_idx=0,
        style_map=style_map, theme_fonts=theme_fonts,
        page_width_cm=page_width_cm,
        left_margin_cm=left_margin_cm,
        right_margin_cm=right_margin_cm,
    )
    paragraphs.extend(table_paras)

    # ── Bước 2: Paragraphs thường (ngoài bảng) ────────────────────────
    # Xây dựng map: numId → abstractNumId để nhận biết các numId thuộc cùng 1 list
    # (Word tạo numId mới khi người dùng ngắt rồi nối tiếp list, nhưng cùng abstractNumId)
    num_to_abstract: dict = {}
    try:
        W_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        wval_ns = f'{{{W_ns}}}val'
        import zipfile as _zf2
        from lxml import etree as _et2
        with _zf2.ZipFile(str(path)) as _z2:
            if 'word/numbering.xml' in _z2.namelist():
                _tree2 = _et2.fromstring(_z2.read('word/numbering.xml'))
                _ns2 = {'w': W_ns}
                for _num_el in _tree2.findall('w:num', _ns2):
                    _nid = _num_el.get(f'{{{W_ns}}}numId')
                    _ref = _num_el.find('w:abstractNumId', _ns2)
                    if _ref is not None:
                        num_to_abstract[_nid] = _ref.get(wval_ns)
    except Exception:
        pass

    # Bộ đếm theo (abstractNumId, ilvl) thay vì (numId, ilvl)
    # → các numId khác nhau nhưng cùng abstractNumId sẽ dùng chung bộ đếm
    abstract_counters: dict = {}

    for para in doc.paragraphs:
        # Bỏ qua para đã xử lý trong bảng ở bước 1
        if id(para._p) in table_para_ids:
            continue

        # ── Xử lý Word List Numbering ─────────────────────────────────
        # Nếu paragraph thuộc numbered list, tính số thứ tự và prefix vào text
        num_id, ilvl = _get_para_numpr(para)
        num_prefix = None
        if num_id is not None:
            key = (num_id, ilvl)
            num_data = numbering_map.get(key, {})
            fmt = num_data.get('fmt', 'decimal')
            if fmt in ('decimal', 'lowerLetter', 'upperLetter',
                       'lowerRoman', 'upperRoman', ''):
                start = num_data.get('start', 1)
                # Dùng (abstractNumId, ilvl) làm key bộ đếm để các numId
                # khác nhau nhưng cùng list (cùng abstractNumId) không reset nhau
                ab_id = num_to_abstract.get(num_id, num_id)
                ab_key = (ab_id, ilvl)
                # Khởi tạo bộ đếm nếu chưa có, hoặc nếu start > counter hiện tại
                # (trường hợp startOverride yêu cầu bắt đầu từ số cụ thể)
                current = abstract_counters.get(ab_key)
                if current is None or start > current + 1:
                    abstract_counters[ab_key] = start - 1
                count = abstract_counters[ab_key] + 1
                abstract_counters[ab_key] = count
                # Giữ tương thích num_counters cũ (dùng trong reset cấp con)
                num_counters[key] = count
                # Reset bộ đếm cấp con khi tăng cấp cha
                for k in list(abstract_counters.keys()):
                    k_ab, k_ilvl = k
                    if k_ab == ab_id and int(k_ilvl) > int(ilvl):
                        # tìm start của cấp con này
                        child_key = next(
                            ((nid, k_ilvl) for nid, ni in num_to_abstract.items()
                             if ni == k_ab),
                            None
                        )
                        child_start = numbering_map.get(child_key, {}).get('start', 1) if child_key else 1
                        abstract_counters[k] = child_start - 1
                # Tạo prefix: "1.", "2.", ...
                num_prefix = f"{count}."

        info = _extract_paragraph(
            para, len(paragraphs), style_map, theme_fonts,
            page_width_cm, left_margin_cm, right_margin_cm,
        )
        if info:
            # Gắn số thứ tự vào text nếu text chưa có sẵn số này
            if num_prefix and not _re_top.match(
                r'^\s*\d+[.)]\s*', info.get('text', '')
            ):
                info['text'] = num_prefix + ' ' + info['text']
                # Cập nhật runs: thêm run số vào đầu
                existing_runs = info.get('runs', [])
                num_run = {'fmt': 'list-number', 'txt': num_prefix + ' '}
                info['runs'] = [num_run] + existing_runs if existing_runs else None
                if info['runs'] is None:
                    del info['runs']
            paragraphs.append(info)

    # ── GIỮ TOÀN BỘ: không cắt vùng chữ ký / nơi nhận ───────────────
    # Tất cả thành phần đều cần kiểm tra thể thức theo NĐ30

    # ── Phát hiện loại văn bản ─────────────────────────────────────────
    para_texts = [p.get('text', '') for p in paragraphs]
    doc_type   = _detect_doc_type(para_texts, filename=path.name)

    # ── Final output ───────────────────────────────────────────────────
    # Thứ tự field: metadata trước, nội dung sau
    result: Dict[str, Any] = {
        'filename': path.name,
        'source':   'docx',     # Tầng 2 biết nguồn gốc file
        'doc_type': doc_type,   # VD: "nghi_dinh", "cong_van", "unknown"
    }

    if page_size:
        result['page_size'] = page_size   # {'width_cm': 21.0, 'height_cm': 29.7}
    if margins:
        result['margins']   = margins     # Chỉ còn lề: top/bottom/left/right (cm)
    if page_number:
        result['page_number'] = page_number
    if paragraphs:
        result['paragraphs']  = paragraphs

    return result


def save_report(report: dict, output_path: str):
    """Lưu report ra file JSON (compact format)"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"✅ Đã lưu: {output_path}")


# ═════════════════════════════════════════════════════════════════════
# CLI RUNNER
# ═════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    import sys
    
    input_dir  = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('data/groundtruth')
    output_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else Path('output')
    output_dir.mkdir(parents=True, exist_ok=True)
    
    docx_files = sorted(input_dir.glob('*.docx'))
    if not docx_files:
        print(f"⚠️  Không tìm thấy file .docx trong: {input_dir}")
        sys.exit(1)
    
    print(f"🔍 Tìm thấy {len(docx_files)} file .docx\n")
    print(f"🎯 Chế độ: OPTIMIZED (Block + Runs, Drop Nulls)\n")
    
    all_reports = []
    
    for docx_path in docx_files:
        print(f"📄 Đang xử lý: {docx_path.name}")
        try:
            report = extract_docx_optimized(str(docx_path))
            
            m = report.get('margins', {})
            pn = report.get('page_number', {})
            print(f"   Lề: trái={m.get('left_cm')}cm | phải={m.get('right_cm')}cm")
            print(f"   Số trang: {pn if pn else 'không phát hiện'}")
            print(f"   Số đoạn văn: {len(report.get('paragraphs', []))}")
            
            # In 3 đoạn đầu để xem cấu trúc
            for p in report.get('paragraphs', [])[:3]:
                text_preview = p['text'][:50]
                block_info = p.get('block', {})
                runs_count = len(p.get('runs', []))
                zone_info = p.get('zone', 'FULL_WIDTH')
                print(f"   ↳ [{p['idx']}] '{text_preview}' | block={block_info} | runs={runs_count} | zone={zone_info}")
            
            out_file = output_dir / f"{docx_path.stem}_optimized.json"
            save_report(report, str(out_file))
            all_reports.append(report)
            print()
            
        except Exception as e:
            import traceback
            print(f"   ❌ Lỗi: {e}")
            traceback.print_exc()
            print()
    
    all_path = output_dir / '_all_optimized.json'
    with open(all_path, 'w', encoding='utf-8') as f:
        json.dump(all_reports, f, ensure_ascii=False, indent=2)
    print(f"\n📦 Tổng hợp toàn bộ: {all_path}")
    print(f"\n🚀 Token giảm ước tính: 70-85% so với bản cũ!")